# Program to create a resume/CV and include Text to speech
# This is a sample

from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# profile picture
document.add_picture('pic.jpg', width=Inches(2.0))

# name, phone number and email details
speak('What is your name? ')
name = input('What is your name? ')
speak('Hello ' + name + 'How are you today?')
speak('What is your phone number?')
phone_number = input('What is your phone number? ')
speak('What is your email address?')
email = input('What is your email address? ')

document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
speak('Tell me about yourself?')
document.add_paragraph(input('Tell me about yourself? '))

#  work experiance
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Enter company mame')
company = input('Enter company mame ')
speak('Your starting date:')
from_date = input('From date: ')
speak('Your leaving Date')
to_date = input('To Date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Describe your experience at ' + company + ' ')
experience_details = input('Describe your experience at ' + company + ' ')
p.add_run(experience_details)

# more experiences
while True:
    speak('Do you have more experiences? Yes or No')
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        speak('Enter company mame')
        company = input('Enter company mame ')
        speak('Your starting date:')
        from_date = input('From date: ')
        speak('Your leaving Date')
        to_date = input('To Date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        speak('Describe your experience at ' + company + ' ')
        experience_details = input(
            'Descibe your experience at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break

# List of Skills
document.add_heading('Skills')
speak('Enter your skill')
skill = input('Enter your skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak('Do you have more skills? Yes or No ')
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        speak('Enter your skill')
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = " Resume generated using Python"


document.save('resume.docx')
